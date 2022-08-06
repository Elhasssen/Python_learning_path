# working with word documents


import os
os.chdir('/home/elhassen/Desktop/Python/learningpy/working_with_pdf')
#######################################
# import docx

# doc = docx.Document('demo.docx')
# print(len(doc.paragraphs))
# # output : 7 
# print(doc.paragraphs[0].text)
# # output : Document Title
# print(doc.paragraphs[1].text)
# # output : A plain paragraph with some bold and some italic
# print(len(doc.paragraphs[1].runs))
# # output : 5 
# # runs are somehow of the styles of the text
# print(doc.paragraphs[1].runs[0].text)
# # output : A plain paragraph with
# # the first one is that 
# print(doc.paragraphs[1].runs[1].text) # output : some
# print(doc.paragraphs[1].runs[2].text) # output : bold
# print(doc.paragraphs[1].runs[3].text) # output : and some

############################################ 
# # getting the full text from a .docx file
# import docx

# def getText(filename):
#     doc = docx.Document(filename)
#     fulltext = []
#     for para in doc.paragraphs:
#         fulltext.append(para.text)
#     return '\n'.join(fulltext)

# ####### 
# # now we read 
# print(getText('demo.docx'))
# ########
# # styling text : 
# # many subfunctions as 
# # if we want to style paragraphs : 
# # ex : doc.paragraphs[0].style = 'Normal'
# # or we can split them into : runs and style one partitions at a time 
# # functions as : .style = "Normal", .underline = True , 
####################################
# writing word documents
import docx

doc = docx.Document()
FirstPar = doc.add_paragraph('Hello, world!')
# if we want to add text to the first paragraphs 
FirstPar.add_run('blab bla this is being added')
# and we can add headings 
doc.add_heading('Big heading', 1)
doc.add_heading('Not that big heading', 2)
doc.add_heading('Tiny big heading', 3)
doc.add_heading('small heading', 4)
# we can add pictures 
doc.add_picture('Image.jpg',width = docx.shared.Inches(1), height = docx.shared.Cm(4))

doc.save('helloworld.docx')

# we can add line and page breaks 
# like ; doc.paragraphs[0].runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)